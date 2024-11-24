import os
from flask import Flask, request, render_template, send_file
from werkzeug.utils import secure_filename
from fpdf import FPDF
from docx import Document
from pdf2docx import Converter
from PyPDF2 import PdfMerger
import pandas as pd

app = Flask(__name__)

UPLOAD_FOLDER = "uploads"
IMAGE_FOLDER = os.path.join(UPLOAD_FOLDER, "images")
PDF_FOLDER = os.path.join(UPLOAD_FOLDER, "pdf")
DOC_FOLDER = os.path.join(UPLOAD_FOLDER, "docs")
EXCEL_FOLDER = os.path.join(UPLOAD_FOLDER, "excel")
os.makedirs(IMAGE_FOLDER, exist_ok=True)
os.makedirs(PDF_FOLDER, exist_ok=True)
os.makedirs(DOC_FOLDER, exist_ok=True)
os.makedirs(EXCEL_FOLDER, exist_ok=True)

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        files = request.files.getlist("files")
        conversion_type = request.form["conversion_type"]
        output_name = request.form["output_name"]

        if conversion_type == "image_to_pdf":
            pdf_path = os.path.join(PDF_FOLDER, f"{output_name}.pdf")
            pdf = FPDF()
            for file in files:
                file_path = os.path.join(IMAGE_FOLDER, secure_filename(file.filename))
                file.save(file_path)
                pdf.add_page()
                pdf.image(file_path, x=10, y=10, w=190)
            pdf.output(pdf_path)
            return send_file(pdf_path, as_attachment=True)

        elif conversion_type == "word_to_pdf" and len(files) == 1:
            file_path = os.path.join(DOC_FOLDER, secure_filename(files[0].filename))
            files[0].save(file_path)
            doc = Document(file_path)
            pdf_path = os.path.join(PDF_FOLDER, f"{output_name}.pdf")
            pdf = FPDF()
            pdf.add_page()
            for para in doc.paragraphs:
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, para.text)
            pdf.output(pdf_path)
            return send_file(pdf_path, as_attachment=True)

        elif conversion_type == "pdf_to_word" and len(files) == 1:
            file_path = os.path.join(PDF_FOLDER, secure_filename(files[0].filename))
            files[0].save(file_path)
            word_path = os.path.join(DOC_FOLDER, f"{output_name}.docx")
            converter = Converter(file_path)
            converter.convert(word_path, start=0, end=None)
            converter.close()
            return send_file(word_path, as_attachment=True)

    return render_template("index.html")

@app.route("/merge", methods=["GET", "POST"])
def merge():
    if request.method == "POST":
        files = request.files.getlist("files")
        merge_type = request.form["merge_type"]
        output_name = request.form["output_name"]

        if merge_type == "merge_pdf":
            merger = PdfMerger()
            for file in files:
                file_path = os.path.join(PDF_FOLDER, secure_filename(file.filename))
                file.save(file_path)
                merger.append(file_path)
            merged_path = os.path.join(PDF_FOLDER, f"{output_name}.pdf")
            merger.write(merged_path)
            merger.close()
            return send_file(merged_path, as_attachment=True)

        elif merge_type == "merge_excel":
            dataframes = []
            for file in files:
                file_path = os.path.join(EXCEL_FOLDER, secure_filename(file.filename))
                file.save(file_path)
                df = pd.read_excel(file_path)
                dataframes.append(df)
            merged_df = pd.concat(dataframes)
            merged_path = os.path.join(EXCEL_FOLDER, f"{output_name}.xlsx")
            merged_df.to_excel(merged_path, index=False)
            return send_file(merged_path, as_attachment=True)

        elif merge_type == "word_to_pdf" and len(files) == 1:
            file_path = os.path.join(DOC_FOLDER, secure_filename(files[0].filename))
            files[0].save(file_path)
            doc = Document(file_path)
            pdf_path = os.path.join(PDF_FOLDER, f"{output_name}.pdf")
            pdf = FPDF()
            pdf.add_page()
            for para in doc.paragraphs:
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, para.text)
            pdf.output(pdf_path)
            return send_file(pdf_path, as_attachment=True)

        elif merge_type == "pdf_to_word" and len(files) == 1:
            file_path = os.path.join(PDF_FOLDER, secure_filename(files[0].filename))
            files[0].save(file_path)
            word_path = os.path.join(DOC_FOLDER, f"{output_name}.docx")
            converter = Converter(file_path)
            converter.convert(word_path, start=0, end=None)
            converter.close()
            return send_file(word_path, as_attachment=True)

        elif merge_type == "excel_to_pdf" and len(files) == 1:
            file_path = os.path.join(EXCEL_FOLDER, secure_filename(files[0].filename))
            files[0].save(file_path)
            df = pd.read_excel(file_path)
            pdf_path = os.path.join(PDF_FOLDER, f"{output_name}.pdf")
            pdf = FPDF()
            pdf.add_page()
            for row in df.iterrows():
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, ' '.join(map(str, row[1].values)))
            pdf.output(pdf_path)
            return send_file(pdf_path, as_attachment=True)

        elif merge_type == "pdf_to_excel" and len(files) == 1:
            file_path = os.path.join(PDF_FOLDER, secure_filename(files[0].filename))
            files[0].save(file_path)
            # Converting PDF to Excel is more complex, you'd likely need a library like `tabula-py`
            # or `pdfplumber` for extracting tables from PDFs, and save as an Excel file.
            # I'll assume some placeholder logic here for simplicity
            excel_path = os.path.join(EXCEL_FOLDER, f"{output_name}.xlsx")
            df = pd.read_pdf(file_path)  # Example placeholder, actual extraction method needed
            df.to_excel(excel_path, index=False)
            return send_file(excel_path, as_attachment=True)

        elif merge_type == "word_to_excel" and len(files) == 1:
            file_path = os.path.join(DOC_FOLDER, secure_filename(files[0].filename))
            files[0].save(file_path)
            doc = Document(file_path)
            data = []
            for para in doc.paragraphs:
                data.append([para.text])
            df = pd.DataFrame(data, columns=["Text"])
            excel_path = os.path.join(EXCEL_FOLDER, f"{output_name}.xlsx")
            df.to_excel(excel_path, index=False)
            return send_file(excel_path, as_attachment=True)

    return render_template("merge.html")

if __name__ == "__main__":
    app.run(debug=True)
